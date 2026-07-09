"""Document loaders for PDF, Word, and Markdown sources."""

import hashlib
from abc import ABC, abstractmethod
from pathlib import Path

from docx import Document as WordDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from enterprise_ai_assistant.models import DocumentSection, LoadedDocument
from enterprise_ai_assistant.rag.exceptions import (
    DocumentLoadError,
    UnsupportedDocumentError,
)


class DocumentLoader(ABC):
    """Base class for one or more supported document suffixes."""

    suffixes: frozenset[str]

    def __init__(self, *, max_file_size_mb: int = 20) -> None:
        """Create a loader with a bounded source file size."""

        if max_file_size_mb <= 0:
            raise ValueError("max_file_size_mb must be positive")
        self._max_file_size_bytes = max_file_size_mb * 1024 * 1024

    def load(self, path: Path) -> LoadedDocument:
        """Validate a path and parse it into logical sections."""

        source = self._validate_path(path)
        try:
            sections = self._load_sections(source)
        except DocumentLoadError:
            raise
        except (OSError, ValueError) as exc:
            raise DocumentLoadError(
                f"failed to load {source.suffix.lower()} document"
            ) from exc
        if not sections:
            raise DocumentLoadError("document contains no extractable text")
        return LoadedDocument(
            id=self._document_id(source),
            source=str(source.absolute()),
            sections=sections,
            metadata={
                "file_name": source.name,
                "file_type": source.suffix.lower().lstrip("."),
                "file_size_bytes": source.stat().st_size,
                "content_sha256": self._content_sha256(source),
            },
        )

    @abstractmethod
    def _load_sections(self, path: Path) -> tuple[DocumentSection, ...]:
        """Parse a validated file into non-empty sections."""

    def _validate_path(self, path: Path) -> Path:
        if path.is_symlink():
            raise DocumentLoadError("symbolic-link documents are not supported")
        if not path.exists():
            raise DocumentLoadError("document path does not exist")
        if not path.is_file():
            raise DocumentLoadError("document path must be a regular file")
        if path.suffix.lower() not in self.suffixes:
            raise UnsupportedDocumentError(
                f"loader does not support {path.suffix.lower()!r}"
            )
        if path.stat().st_size > self._max_file_size_bytes:
            raise DocumentLoadError("document exceeds the configured size limit")
        return path

    @staticmethod
    def _document_id(path: Path) -> str:
        # The source path is stable across content revisions, so re-ingestion
        # replaces old chunks instead of creating a second logical document.
        source_key = str(path.absolute()).encode("utf-8")
        return hashlib.sha256(source_key).hexdigest()[:24]

    @staticmethod
    def _content_sha256(path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as source:
            while block := source.read(1024 * 1024):
                digest.update(block)
        return digest.hexdigest()


class MarkdownLoader(DocumentLoader):
    """Load UTF-8 Markdown while preserving headings and list markers."""

    suffixes = frozenset({".md", ".markdown"})

    def _load_sections(self, path: Path) -> tuple[DocumentSection, ...]:
        try:
            content = path.read_text(encoding="utf-8")
        except UnicodeDecodeError as exc:
            raise DocumentLoadError("Markdown document must use UTF-8") from exc
        if not content.strip():
            return ()
        return (
            DocumentSection(
                content=content,
                metadata={"content_type": "markdown"},
            ),
        )


class WordLoader(DocumentLoader):
    """Load DOCX paragraphs and tables without requiring Microsoft Word."""

    suffixes = frozenset({".docx"})

    def _load_sections(self, path: Path) -> tuple[DocumentSection, ...]:
        document = WordDocument(str(path))
        sections: list[DocumentSection] = []

        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        if paragraphs:
            sections.append(
                DocumentSection(
                    content="\n\n".join(paragraphs),
                    metadata={"content_type": "body"},
                )
            )

        for table_index, table in enumerate(document.tables):
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells)
                for row in table.rows
            ]
            content = "\n".join(row for row in rows if row.strip(" |"))
            if content.strip():
                sections.append(
                    DocumentSection(
                        content=content,
                        metadata={
                            "content_type": "table",
                            "table_index": table_index,
                        },
                    )
                )
        return tuple(sections)


class PDFLoader(DocumentLoader):
    """Load text PDF pages and retain one-based page numbers."""

    suffixes = frozenset({".pdf"})

    def _load_sections(self, path: Path) -> tuple[DocumentSection, ...]:
        try:
            reader = PdfReader(str(path))
        except PdfReadError as exc:
            raise DocumentLoadError("PDF structure is invalid") from exc
        if reader.is_encrypted:
            raise DocumentLoadError("encrypted PDF documents are not supported")

        sections: list[DocumentSection] = []
        for page_number, page in enumerate(reader.pages, start=1):
            # Layout mode preserves columns and line boundaries better than
            # plain extraction, which improves downstream semantic chunking.
            content = page.extract_text(extraction_mode="layout") or ""
            if content.strip():
                sections.append(
                    DocumentSection(
                        content=content,
                        metadata={
                            "content_type": "page",
                            "page_number": page_number,
                        },
                    )
                )
        return tuple(sections)


class DocumentLoaderRegistry:
    """Resolve document loaders by normalized file suffix."""

    def __init__(self, loaders: tuple[DocumentLoader, ...]) -> None:
        """Create a registry and reject ambiguous suffix ownership."""

        self._loaders: dict[str, DocumentLoader] = {}
        for loader in loaders:
            for suffix in loader.suffixes:
                if suffix in self._loaders:
                    raise ValueError(f"duplicate document loader for {suffix}")
                self._loaders[suffix] = loader

    def load(self, path: Path) -> LoadedDocument:
        """Load a document using the registered suffix adapter."""

        try:
            loader = self._loaders[path.suffix.lower()]
        except KeyError as exc:
            raise UnsupportedDocumentError(
                f"unsupported document suffix: {path.suffix.lower()!r}"
            ) from exc
        return loader.load(path)


def create_document_loader_registry(
    *,
    max_file_size_mb: int = 20,
) -> DocumentLoaderRegistry:
    """Build the default PDF, DOCX, and Markdown loader registry."""

    return DocumentLoaderRegistry(
        (
            PDFLoader(max_file_size_mb=max_file_size_mb),
            WordLoader(max_file_size_mb=max_file_size_mb),
            MarkdownLoader(max_file_size_mb=max_file_size_mb),
        )
    )

"""Tests for PDF, Word, Markdown, and loader dispatch."""

from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document as WordDocument

from enterprise_ai_assistant.rag import (
    DocumentLoadError,
    MarkdownLoader,
    PDFLoader,
    UnsupportedDocumentError,
    WordLoader,
    create_document_loader_registry,
)
from enterprise_ai_assistant.rag import loaders as loader_module


def test_markdown_loader_preserves_content_and_metadata(tmp_path: Path) -> None:
    """Markdown should remain intact for heading-aware chunk boundaries."""

    path = tmp_path / "guide.md"
    path.write_text("# 标题\n\n企业知识内容", encoding="utf-8")

    document = MarkdownLoader().load(path)

    assert document.sections[0].content.startswith("# 标题")
    assert document.sections[0].metadata["content_type"] == "markdown"
    assert document.metadata["file_type"] == "md"
    assert len(document.metadata["content_sha256"]) == 64  # type: ignore[arg-type]


def test_word_loader_reads_paragraphs_and_tables(tmp_path: Path) -> None:
    """DOCX body paragraphs and tables should become inspectable sections."""

    path = tmp_path / "manual.docx"
    source = WordDocument()
    source.add_heading("产品手册", level=1)
    source.add_paragraph("退款将在三个工作日内到账。")
    table = source.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "套餐"
    table.cell(0, 1).text = "企业版"
    source.save(path)

    document = WordLoader().load(path)

    assert len(document.sections) == 2
    assert "退款将在三个工作日内到账" in document.sections[0].content
    assert document.sections[1].content == "套餐 | 企业版"
    assert document.sections[1].metadata["table_index"] == 0


def test_pdf_loader_keeps_page_numbers(
    tmp_path: Path,
    monkeypatch: object,
) -> None:
    """Extracted PDF pages should retain one-based page metadata."""

    path = tmp_path / "policy.pdf"
    path.write_bytes(b"%PDF-fake")
    pages = [
        SimpleNamespace(extract_text=lambda **kwargs: "第一页制度"),
        SimpleNamespace(extract_text=lambda **kwargs: ""),
        SimpleNamespace(extract_text=lambda **kwargs: "第三页流程"),
    ]
    fake_reader = SimpleNamespace(is_encrypted=False, pages=pages)
    monkeypatch.setattr(  # type: ignore[attr-defined]
        loader_module,
        "PdfReader",
        lambda _: fake_reader,
    )

    document = PDFLoader().load(path)

    assert [section.metadata["page_number"] for section in document.sections] == [
        1,
        3,
    ]


def test_loader_registry_rejects_unsupported_suffix(tmp_path: Path) -> None:
    """Unsupported binary formats should fail before content inspection."""

    path = tmp_path / "legacy.doc"
    path.write_bytes(b"legacy")
    registry = create_document_loader_registry()

    with pytest.raises(UnsupportedDocumentError, match="unsupported"):
        registry.load(path)


def test_loader_rejects_oversized_file(tmp_path: Path) -> None:
    """Configured file-size limits should be enforced before parsing."""

    path = tmp_path / "large.md"
    path.write_bytes(b"x" * (1024 * 1024 + 1))
    loader = MarkdownLoader(max_file_size_mb=1)

    with pytest.raises(DocumentLoadError, match="size limit"):
        loader.load(path)


def test_loader_rejects_empty_document(tmp_path: Path) -> None:
    """Whitespace-only sources should not enter the ingestion pipeline."""

    path = tmp_path / "empty.md"
    path.write_text("  \n", encoding="utf-8")

    with pytest.raises(DocumentLoadError, match="no extractable text"):
        MarkdownLoader().load(path)
